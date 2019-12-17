#!/usr/bin/env python
# coding: utf-8
import glob
from docx import Document
from docx.shared import Inches, Pt
import cv2
import numpy as np
import docx
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
import make_read_qr
QNUM = 20

# Wordファイルのオブジェクトを生成
document = Document()

# 文書を横向き(landscape)に変更
# section = document.sections[-1]
section = document.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
section.orientation = WD_ORIENT.LANDSCAPE
# print(section.start_type)
# print(section.orientation)
section.right_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section1 = document.sections[-1]
section1.orientation = WD_ORIENT.LANDSCAPE

paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.space_before = Pt(0)
paragraph_format.first_line_indent = Pt(0)
paragraph_format.right_indent = Pt(0)

document.add_heading('解答用紙          名前 :', 0).add_run().font
# font = document.add_paragraph().add_run().font
# font.name = 'Calibri'
# font.size = Pt(12)
# 問題識別用のQRコード
hash_qr_paths = glob.glob('./qr/qr_storage/hash*')
document.add_picture(hash_qr_paths[0])
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
# docimg0 = document.add_picture('./img/gray_img.jpg', width=Inches(1.25))
# temp = document.add_paragraph().add_run().add_picture('./img/final22.jpg')

# img_paragraph = document.add_paragraph()
# docimg =document.add_picture('./img/gray_img.jpg', width=Inches(1.25))
# img = cv2.imread('./img/gray_img.jpg')
print('done!!')

img = cv2.imread('./qr/qr_storage/qr_1.png')
height, width = img.shape[:2]
white_img = 255 - np.zeros((height, width, 1), np.uint8)
cv2.imwrite('./qr/qr_storage/white_img.png', white_img)
table = document.add_table(rows=0, cols=4)
print(height, width)
qno = 0
flag = False
for i in range(0,9):
    if qno + 3 >= QNUM:
        break
    row_cells = table.add_row().cells
    table.rows.height = 1000
    # row_cells[0].text = str(i+1)
    for j in range(0,4):
        if qno == QNUM:
            flag = True
            break
        qno += 1
        row_cells[j].text = str(qno)
        qr_table = row_cells[j].add_table(rows=1, cols=2)
        qr_header = qr_table.columns[0].cells[0].add_paragraph().add_run().add_picture('./qr/qr_storage/white_img.png', width=Inches(0.7))
        qr_table.columns[1].cells[0].add_paragraph().add_run().add_picture('./qr/qr_storage/qr_{}.png'.format(qno), width=Inches(0.7))
    if flag:
        break

document.save('testsheet.docx')